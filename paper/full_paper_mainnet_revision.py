#!/usr/bin/env python3
"""
Full-paper revision: replace simulated metrics with mainnet-measured data.
Revised text is highlighted (blue font + yellow background) in the output DOCX.
"""

import json
import re
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
SRC = ROOT / 'TENDERFLOW_Paper_Mainnet_Data_20260603.docx'
EXP = ROOT / 'raw-data' / 'experiment_latest.json'
OUT = ROOT / 'TENDERFLOW_Paper_Mainnet_Revised_Highlighted_20260603.docx'
CHANGELOG = ROOT / 'REVISION_CHANGELOG_20260603.md'

REVISED_RGB = RGBColor(0, 51, 153)  # dark blue for revised text


def load_metrics():
    with open(EXP, encoding='utf-8') as f:
        exp = json.load(f)
    sc = exp['scenario_analysis']
    bench = exp['benchmark']
    sim = exp['simulation']
    net = exp['network']
    txs = sim['transactions']
    ok = sum(1 for t in txs if t.get('ok'))
    total = len(txs)
    return {
        'block_height': net['block_height'],
        'nodes': net['infrastructure_nodes'],
        'orderers': net['orderer_count'],
        'peers': net['peer_count'],
        'fabric': net['fabric_version'],
        'channel': net['channel_id'],
        'tps': bench['tps'],
        'p95_ms': bench['latency_ms']['p95'],
        'avg_ms': bench['latency_ms']['avg'],
        'sim_txs': total,
        'sim_ok': ok,
        'sim_success_pct': round(100 * ok / total, 1) if total else 0,
        'blocks_sim': sim['blocks_produced'],
        'static_winner': sc['scenario1_static_winner'],
        'static_vi': sc['scenario1_static_verified_rep'],
        'dynamic_winner': sc['scenario2_dynamic_winner'],
        'dynamic_score': round(sc['scenario2_dynamic_total_score'], 3),
        'bidder_b_total': round(sc['bidder_b_total_after_breach'], 3),
        'bidder_c_score': round(sc['all_scores']['Bidder-C']['total_score'], 3),
        'experiment_id': exp['experiment_id'],
    }


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def append_run(paragraph, text, revised=False):
    if not text:
        return
    run = paragraph.add_run(text)
    if revised:
        run.font.color.rgb = REVISED_RGB
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.font.bold = False


def apply_replacements_colored(paragraph, replacements):
    """replacements: list of (old, new) applied in order; new segments colored."""
    text = paragraph.text
    if not text.strip():
        return []

    changes = []
    for old, new in replacements:
        if old in text:
            text = text.replace(old, new)
            changes.append(f'"{old[:60]}..." → "{new[:60]}..."' if len(old) > 60 else f'"{old}" → "{new}"')

    if text == paragraph.text:
        return []

    clear_paragraph(paragraph)
    # Rebuild with coloring on replaced segments
    remaining = paragraph.text if False else text
    # Walk original with markers
    orig = paragraph.text
    # Simpler: split using first replacement set combined
    work = paragraph.text
    # Re-read from stored original before clear
    return changes


def set_paragraph_from_parts(paragraph, parts):
    """parts: list of (text, is_revised)"""
    clear_paragraph(paragraph)
    for text, revised in parts:
        append_run(paragraph, text, revised=revised)


def replace_once_colored(paragraph, old, new, change_log):
    full = paragraph.text
    if old not in full:
        return False
    idx = full.find(old)
    before = full[:idx]
    after = full[idx + len(old):]
    set_paragraph_from_parts(paragraph, [
        (before, False),
        (new, True),
        (after, False),
    ])
    change_log.append({'type': 'replace', 'old': old, 'new': new})
    return True


def replace_all_in_paragraph(paragraph, replacements, change_log):
    """Apply multiple replacements sequentially, coloring each new insert."""
    text = paragraph.text
    if not text.strip():
        return
    applied = False
    for old, new in replacements:
        if old not in text:
            continue
        applied = True
        change_log.append({'type': 'replace', 'old': old, 'new': new})
        text = text.replace(old, new)
    if not applied:
        return
    # Color entire paragraph if any replacement (simpler for multi-replace)
    clear_paragraph(paragraph)
    append_run(paragraph, text, revised=True)


def iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def add_revision_front_matter(doc, m, change_count):
    # Insert after title block — find Abstract:
    abstract_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().lower() == 'abstract:')

    note_lines = [
        '【修订说明 / Revision Note — Mainnet Data Update】',
        f'修订日期 Revision date: {datetime.now(timezone.utc).strftime("%Y-%m-%d")} UTC',
        f'实验编号 Experiment ID: {m["experiment_id"]}',
        '蓝色字体 + 黄色底纹 Blue text on yellow highlight = 本次依据主网实测数据修订的内容 '
        '(Revised passages based on live 14-node Fabric mainnet measurements).',
        f'主网摘要 Mainnet snapshot: channel={m["channel"]}, block height={m["block_height"]}, '
        f'{m["nodes"]} nodes (Fabric {m["fabric"]}), {m["sim_txs"]} simulation invokes '
        f'({m["sim_success_pct"]}% success), measured TPS={m["tps"]}, p95 latency={m["p95_ms"]} ms.',
        f'共 {change_count} 处文本修订 {change_count} text revisions applied. '
        f'详见 paper/REVISION_CHANGELOG_20260603.md',
    ]
    for line in reversed(note_lines):
        new_p = doc.paragraphs[abstract_idx]._element
        p_el = doc.paragraphs[abstract_idx]._element.makeelement(qn('w:p'), {})
        new_p.addprevious(p_el)
        from docx.text.paragraph import Paragraph
        para = Paragraph(p_el, doc.paragraphs[abstract_idx]._parent)
        run = para.add_run(line)
        run.font.color.rgb = REVISED_RGB
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.font.size = Pt(10)
        if line.startswith('【'):
            run.bold = True


def build_replacements(m):
    abstract_old = (
        "Experimental simulation of the Ma'anshan Yangtze River Rail cum Road Bridge demonstrates "
        "that the proposed system achieves a peak throughput of 350 transactions per second (TPS) "
        "and reduces verification latency from days to seconds. Furthermore, our quantitative "
        "analysis reveals a 71% decrease in annual administrative and audit costs per enterprise "
        "by replacing manual oversight with cryptographic proof."
    )
    abstract_new = (
        f"A live {m['nodes']}-node Hyperledger Fabric mainnet (five RAFT orderers, nine endorsing peers, "
        f"CouchDB state, CCAAS tenderflow chaincode on channel {m['channel']}) executed a Ma'anshan "
        f"Yangtze River Bridge tendering simulation of {m['sim_txs']} on-chain invokes "
        f"(CreateTender, CommitBid, RevealBid, ReputationUpdate), producing {m['block_height']} ledger "
        f"blocks with {m['sim_success_pct']}% transaction success ({m['sim_ok']}/{m['sim_txs']}). "
        f"Measured sequential end-to-end invoke throughput was {m['tps']} TPS with p95 latency "
        f"{m['p95_ms']} ms (average {m['avg_ms']} ms). The on-chain reputation engine filtered "
        f"Bidder B after a simulated breach reveal; {m['dynamic_winner']} achieved the highest "
        f"composite score ({m['dynamic_score']}). These mainnet-measured outcomes validate dynamic "
        f"on-chain behavior reputation for high-integrity procurement governance."
    )

    return [
        (abstract_old, abstract_new),
        (
            'peak throughput of 350 transactions per second (TPS)',
            f'measured sequential throughput of {m["tps"]} transactions per second (TPS) on the deployed mainnet',
        ),
        (
            'peak throughput of 350 TPS',
            f'measured sequential throughput of {m["tps"]} TPS on the 14-node mainnet',
        ),
        (
            'The network achieved a peak throughput of 350 TPS under the Raft consensus protocol, which is more than sufficient for the concurrent submission of complex technical bids and bid bonds.',
            f'On the deployed 14-node RAFT mainnet (Fabric {m["fabric"]}), sequential invoke benchmarking '
            f'recorded {m["tps"]} TPS across 30 CommitBid transactions with p95 latency {m["p95_ms"]} ms '
            f'(average {m["avg_ms"]} ms), sufficient for sealed-bid phases of mega-project tendering where '
            f'endorsement rounds—not unbounded fan-in concurrency—dominate latency.',
        ),
        (
            'Even under high load scenarios (1,000 concurrent requests), the latency for the "Commit" transaction remained below 150ms, providing a seamless user experience for bidding consortiums.',
            f'Measured mainnet CommitBid invoke latency averaged {m["avg_ms"]} ms with p95 '
            f'{m["p95_ms"]} ms under sequential load (30 invokes), reflecting realistic multi-org '
            f'endorsement and ordering on the Ma\'anshan simulation channel.',
        ),
        (
            'Traditional manual verification of contractor qualifications and bid bond authenticity typically takes 45 days; the TENDERFLOW system reduced this to under 2 seconds via automated smart contract validation.',
            f'On-chain RevealBid integrity checks and GetReputation queries completed within the measured '
            f'invoke/query window (p95 {m["p95_ms"]} ms), replacing manual multi-week qualification '
            f'reviews with deterministic chaincode validation in the mainnet experiment.',
        ),
        (
            'By replacing centralized third-party clearinghouses with cryptographic proofs, the simulated annual administrative and audit costs per enterprise dropped from $75,000 to $22,000, representing a 71% net reduction.',
            f'Direct mainnet cost accounting was outside the scope of this deployment; administrative '
            f'efficiency is evidenced by {m["sim_txs"]} automated lifecycle invokes with '
            f'{m["sim_success_pct"]}% success and {m["blocks_sim"]} blocks produced without manual ledger reconciliation.',
        ),
        (
            'The simulation results indicate that the TENDERFLOW mode enhances the Integrity of the Selection Pool by 40% compared to traditional modes.',
            f'Mainnet simulation showed measurable integrity gain: Bidder B (initial Vi=0.95) was filtered '
            f'after a failed reveal (total score {m["bidder_b_total"]}), whereas static qualification alone '
            f'would have retained equal-ranked bidders A and B.',
        ),
        (
            'The contract was awarded to Bidder C, who possessed a lower  but a significantly higher due to 12 consecutive successful "Commit and Reveal" executions.',
            f'{m["dynamic_winner"]} ranked highest on-chain (total score {m["dynamic_score"]}) after Bidder B '
            f'was penalized; Bidder C rose to {m["bidder_c_score"]} following 12 consecutive Success reputation updates.',
        ),
        (
            'The contract was awarded to Bidder C, who possessed a lower Vi but a significantly higher behavioral score due to 12 consecutive successful "Commit and Reveal" executions.',
            f'{m["dynamic_winner"]} ranked highest on-chain (total score {m["dynamic_score"]}) after Bidder B '
            f'was penalized; Bidder C reached {m["bidder_c_score"]} after 12 on-chain Success updates.',
        ),
        (
            'With a peak throughput of 350 TPS and a block finality of 2 seconds, the system ensures that critical bidding data is anchored nearly instantaneously.',
            f'With measured sequential throughput of {m["tps"]} TPS, p95 latency {m["p95_ms"]} ms, and '
            f'{m["block_height"]} blocks on the live mainnet, bidding payloads were anchored with '
            f'multi-org RAFT ordering during the Ma\'anshan simulation.',
        ),
        (
            'The informatics-driven approach achieved a 99.9% reduction in verification latency, transforming a 45-day manual process into a sub-2-second automated check. Furthermore, replacing centralized oversight with cryptographic proof led to a 71% decrease in administrative and audit costs.',
            f'Mainnet invokes achieved {m["sim_success_pct"]}% success across {m["sim_txs"]} lifecycle '
            f'transactions with p95 latency {m["p95_ms"]} ms, replacing manual qualification review with '
            f'on-chain RevealBid and reputation updates; formal cost-benefit analysis remains future work.',
        ),
        (
            'While 350 TPS is sufficient for individual bridge projects, industry-wide adoption across thousands of concurrent tenders may require advanced sharding techniques or Layer-2 scaling solutions.',
            f'While {m["tps"]} TPS (sequential mainnet benchmark) suffices for sealed-bid phases of '
            f'mega-projects, industry-wide concurrent tender peaks may require horizontal peer scaling '
            f'or batching optimizations beyond this single-channel deployment.',
        ),
        (
            'research into Layer-2 scaling solutions (like state channels) is necessary to increase throughput beyond 350 TPS',
            f'horizontal scaling and invoke batching should be studied to raise throughput beyond the '
            f'measured {m["tps"]} TPS sequential baseline',
        ),
        (
            'Automated verification of bid bonds and credentials reduces the verification latency from 45 days to under 2 seconds, drastically lowering the administrative overhead and human error associated with manual auditing.',
            f'Automated on-chain RevealBid and reputation queries in the mainnet test completed within '
            f'{m["p95_ms"]} ms (p95), reducing manual reconciliation during the Ma\'anshan simulation.',
        ),
        (
            'Despite the documented efficiency gains (71% cost reduction), several challenges remain for large-scale deployment:',
            'Despite documented mainnet performance and integrity gains, several challenges remain for large-scale deployment:',
        ),
        ('350 TPS', f'{m["tps"]} TPS (mainnet)'),
        (
            ' The results validate that transitioning from static qualification to dynamic onchain behavior reputation effectively bridges the gap between procurement ideals and practice, offering a strategic blueprint for high integrity digital governance in the construction sector.',
            '',
        ),
        (
            'Transaction Throughput (TPS): The network achieved a measured sequential throughput of 2.14 TPS on the 14-node mainnet under the Raft consensus protocol, which is more than sufficient for the concurrent submission of complex technical bids and bid bonds.',
            f'Transaction Throughput (TPS): Sequential mainnet benchmarking recorded {m["tps"]} TPS '
            f'(30 CommitBid invokes) on the 14-node RAFT deployment (Fabric {m["fabric"]}), with p95 latency '
            f'{m["p95_ms"]} ms—adequate for sealed-bid tendering dominated by multi-org endorsement.',
        ),
        ('71% decrease', 'mainnet-measured automation (cost study pending)'),
        ('71% net reduction', 'automated on-chain lifecycle (cost study pending)'),
        ('71% cost reduction', 'mainnet invoke automation (cost study pending)'),
        ('71% decrease in administrative', 'reduction of manual ledger steps (71% cost figure removed— not measured on mainnet)'),
    ]


def write_changelog(log, m):
    lines = [
        '# TENDERFLOW Paper — Mainnet Revision Changelog',
        '',
        f'**Output file:** `{OUT.name}`',
        f'**Experiment:** `{m["experiment_id"]}`',
        f'**Source data:** `paper/raw-data/experiment_latest.json`',
        '',
        '## Highlight legend',
        '- **Blue font + yellow highlight** in DOCX = text revised to reflect live mainnet measurements.',
        '',
        '## Mainnet metrics used',
        f'- Channel: `{m["channel"]}`',
        f'- Block height: **{m["block_height"]}**',
        f'- Infrastructure: **{m["nodes"]}** nodes ({m["orderers"]} orderers + {m["peers"]} peers)',
        f'- Fabric: **{m["fabric"]}** (CCAAS chaincode)',
        f'- Simulation invokes: **{m["sim_txs"]}** ({m["sim_ok"]} OK, {m["sim_success_pct"]}%)',
        f'- Blocks produced (simulation phase): **{m["blocks_sim"]}**',
        f'- Throughput: **{m["tps"]} TPS** (30 sequential CommitBid benchmark)',
        f'- Latency p95 / avg: **{m["p95_ms"]} ms / {m["avg_ms"]} ms**',
        f'- Static winner: **{m["static_winner"]}** (Vi={m["static_vi"]})',
        f'- Dynamic winner: **{m["dynamic_winner"]}** (total={m["dynamic_score"]})',
        f'- Bidder B after breach: total **{m["bidder_b_total"]}** (filtered)',
        f'- Bidder C after 12 updates: total **{m["bidder_c_score"]}**',
        '',
        '## Removed or qualified (not measured on mainnet)',
        '- **350 TPS** — replaced with measured **2.14 TPS** sequential benchmark',
        '- **71% administrative cost reduction** — removed (no cost instrumentation on mainnet)',
        '- **45 days → 2 seconds** — replaced with measured **p95 416 ms** invoke latency',
        '- **+40% selection pool integrity** — replaced with Bidder B breach filtering evidence',
        '- **Bidder C awarded** — corrected to **Bidder A** highest total score on mainnet',
        '',
        '## Text replacements applied',
        '',
    ]
    for i, entry in enumerate(log, 1):
        old = entry['old'].replace('\n', ' ')
        new = entry['new'].replace('\n', ' ')
        lines.append(f'{i}. `{old[:100]}{"..." if len(old)>100 else ""}`')
        lines.append(f'   → `{new[:100]}{"..." if len(new)>100 else ""}`')
        lines.append('')
    CHANGELOG.write_text('\n'.join(lines), encoding='utf-8')


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    m = load_metrics()
    doc = Document(str(SRC))
    replacements = build_replacements(m)
    change_log = []

    for para in iter_all_paragraphs(doc):
        for old, new in replacements:
            if old in para.text:
                replace_once_colored(para, old, new, change_log)
                break  # re-scan paragraph after structural change

    # Second pass for any remaining shorter patterns
    for para in iter_all_paragraphs(doc):
        for old, new in replacements:
            if old in para.text and old != new:
                replace_once_colored(para, old, new, change_log)

    add_revision_front_matter(doc, m, len(change_log))
    doc.save(str(OUT))
    write_changelog(change_log, m)
    print(f'Saved: {OUT}')
    print(f'Changelog: {CHANGELOG}')
    print(f'Revisions: {len(change_log)}')


if __name__ == '__main__':
    main()
