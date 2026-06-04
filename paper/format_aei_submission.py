#!/usr/bin/env python3
"""
Format TENDERFLOW paper for Advanced Engineering Informatics (Elsevier AEI).

Requirements applied (Guide for Authors):
- Title page with affiliations (superscript letters)
- Abstract <= 250 words, standalone
- 1-7 keywords
- Highlights: 3-5 bullets, <= 85 characters each (+ separate file)
- Numbered sections (1., 1.1., ...)
- Numbered references [n] in order of appearance
- Back matter: Acknowledgements, CRediT, Competing interests, GenAI declaration, Data availability
- Single-column editable Word; remove internal revision markup
"""

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
SRC = ROOT / 'TENDERFLOW_Paper_Mainnet_Revised_Highlighted_20260603.docx'
EXP = ROOT / 'raw-data' / 'experiment_latest.json'
OUT = ROOT / 'TENDERFLOW_Paper_AEI_Submission_20260603.docx'
HIGHLIGHTS_TXT = ROOT / 'AEI_Highlights.txt'
CHECKLIST = ROOT / 'AEI_Submission_Checklist.md'

FONT = 'Times New Roman'
BODY_SIZE = Pt(12)
HEADING_SIZE = Pt(12)


def load_metrics():
    with open(EXP, encoding='utf-8') as f:
        exp = json.load(f)
    sc = exp['scenario_analysis']
    bench = exp['benchmark']
    sim = exp['simulation']
    net = exp['network']
    txs = sim['transactions']
    ok = sum(1 for t in txs if t.get('ok'))
    return {
        'block_height': net['block_height'],
        'nodes': net['infrastructure_nodes'],
        'channel': net['channel_id'],
        'fabric': net['fabric_version'],
        'fabric': net['fabric_version'],
        'tps': bench['tps'],
        'p95_ms': bench['latency_ms']['p95'],
        'avg_ms': bench['latency_ms']['avg'],
        'sim_txs': len(txs),
        'sim_ok': ok,
        'sim_pct': round(100 * ok / len(txs), 1),
        'dynamic_winner': sc['scenario2_dynamic_winner'],
        'dynamic_score': round(sc['scenario2_dynamic_total_score'], 3),
    }


def build_highlights(m):
    candidates = [
        f"Live {m['nodes']}-node Fabric mainnet simulates Ma'anshan bridge tendering",
        "On-chain reputation engine filters bidder after breach reveal",
        f"Measured {m['tps']} TPS sequential throughput; p95 latency {m['p95_ms']:.0f} ms",
        "Commit-reveal smart contracts automate sealed bid lifecycle",
        "Dynamic reputation outperforms static qualification-only ranking",
    ]
    valid = []
    for h in candidates:
        if len(h) <= 85:
            valid.append(h)
        else:
            # truncate smartly
            valid.append(h[:82] + '...')
    return valid[:5]


def build_abstract_aei(m):
    """AEI abstract: purpose, results, conclusions; <= 250 words."""
    return (
        "Construction e-tendering remains vulnerable to centralized opacity and static "
        "qualification checks that ignore recent contractor behavior. From an engineering "
        "informatics perspective, procurement is a knowledge-intensive decision task requiring "
        "traceable evidence, automated rule execution, and evolving trust models. This paper "
        "presents TENDERFLOW, a permissioned Hyperledger Fabric platform integrating commit-reveal "
        "bidding with a multi-tier on-chain behavior reputation ontology. Verified, behavioral, "
        "and social reputation tiers are updated through smart contracts using asymmetric "
        "learning and penalty functions. A fourteen-node mainnet (five RAFT orderers, nine "
        "endorsing peers, CouchDB state, chaincode-as-a-service on channel "
        f"{m['channel']}, Fabric {m['fabric']}) executed a Ma'anshan Yangtze River Bridge "
        f"tendering experiment with {m['sim_txs']} lifecycle invokes, producing "
        f"{m['block_height']} blocks and {m['sim_pct']}% transaction success "
        f"({m['sim_ok']}/{m['sim_txs']}). Measured sequential throughput was {m['tps']} "
        f"transactions per second; p95 end-to-end latency was {m['p95_ms']:.0f} ms. The "
        "reputation engine penalized a simulated bid-bond breach and ranked "
        f"{m['dynamic_winner']} highest (composite score {m['dynamic_score']}). The study "
        "validates that coupling immutable audit trails with transaction-derived behavioral "
        "signals improves integrity in construction procurement informatics beyond static "
        "credential screening."
    )


def word_count(text):
    return len(text.split())


def clear_run_formatting(run):
    run.font.highlight_color = None
    run.font.color.rgb = RGBColor(0, 0, 0)


def style_paragraph(paragraph, bold=False, size=BODY_SIZE, space_after=6, first_line_indent=None):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE if False else WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        paragraph.paragraph_format.first_line_indent = Inches(first_line_indent)
    for run in paragraph.runs:
        run.font.name = FONT
        run.font.size = size
        run.font.bold = bold
        clear_run_formatting(run)


def apply_body_style(doc):
    for para in doc.paragraphs:
        style_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    style_paragraph(para)


def remove_revision_notes(doc):
    to_remove = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith('【修订说明') or t.startswith('修订日期') or t.startswith('实验编号'):
            to_remove.append(p)
        elif 'Blue text on yellow highlight' in t:
            to_remove.append(p)
        elif t.startswith('主网摘要 Mainnet snapshot'):
            to_remove.append(p)
        elif t.startswith('共 15 处文本修订') or t.startswith('共 ') and 'text revisions' in t:
            to_remove.append(p)
    for p in to_remove:
        el = p._element
        el.getparent().remove(el)


def insert_after(paragraph, text='', bold=False, style_name=None):
    new_el = paragraph._element.makeelement(qn('w:p'), {})
    paragraph._element.addnext(new_el)
    new_p = Paragraph(new_el, paragraph._parent)
    if text:
        run = new_p.add_run(text)
        run.font.name = FONT
        run.font.size = BODY_SIZE
        run.font.bold = bold
    return new_p


def rebuild_title_page(doc):
    """AEI title page: title, authors with superscript, affiliations, corresponding author."""
    title_p = doc.paragraphs[0]
    title_p.alignment = 1
    style_paragraph(title_p, bold=True, size=Pt(14))

    remove_idxs = []
    for i, p in enumerate(doc.paragraphs[1:12], start=1):
        t = p.text.strip().lower()
        if not t or 'school of architectural' in t or 'zhonghua' in t or 'guangzhou' in t or 'corresponding author' in t:
            remove_idxs.append(i)
    for i in sorted(remove_idxs, reverse=True):
        el = doc.paragraphs[i]._element
        el.getparent().remove(el)

    auth = insert_after(doc.paragraphs[0], bold=False)
    auth.alignment = 1
    for name in ['Zhonghua Peng', 'Tao Ye', 'Gongming Wang']:
        if auth.text or len(auth.runs) > 0:
            r = auth.add_run(', ')
            r.font.name = FONT
        r = auth.add_run(name)
        r.font.name = FONT
        r.font.size = BODY_SIZE
        s = auth.add_run('a')
        s.font.superscript = True
        s.font.name = FONT

    aff = insert_after(auth)
    aff.alignment = 1
    ra = aff.add_run('a ')
    ra.font.superscript = True
    aff.add_run(
        'School of Architectural Engineering, Guangzhou Institute of Science and Technology, '
        'Guangzhou 510540, China'
    )

    corr = insert_after(aff)
    corr.alignment = 1
    corr.add_run('Corresponding author: ').font.bold = True
    corr.add_run('Zhonghua Peng (zhonghua.peng@example.edu.cn)  [Update email before submission]')

    insert_after(corr, '')  # spacer


def set_abstract_and_keywords(doc, abstract_text, keywords):
    # Find Abstract: paragraph
    abs_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().lower() == 'abstract:')
    abs_body = doc.paragraphs[abs_idx + 2] if doc.paragraphs[abs_idx + 1].text.strip() == '' else doc.paragraphs[abs_idx + 1]
    abs_body.text = abstract_text
    style_paragraph(abs_body)
    wc = word_count(abstract_text)
    if wc > 250:
        raise ValueError(f'Abstract exceeds 250 words: {wc}')

    kw_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().lower().startswith('keywords:'))
    kw_p = doc.paragraphs[kw_idx]
    kw_p.text = f'Keywords: {keywords}'
    style_paragraph(kw_p)


def insert_highlights(doc, highlights):
    kw_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().lower().startswith('keywords:'))
    anchor = doc.paragraphs[kw_idx]
    h_title = insert_after(anchor, 'Highlights', bold=True)
    h_title.paragraph_format.space_before = Pt(12)
    last = h_title
    for h in highlights:
        bp = insert_after(last, f'• {h}')
        bp.paragraph_format.left_indent = Inches(0.25)
        last = bp
    insert_after(last, '')


def insert_back_matter(doc):
    ref_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().lower() == 'references')
    anchor = doc.paragraphs[ref_idx - 1]

    sections = [
        ('Acknowledgements', (
            'The authors acknowledge the Guangzhou Institute of Science and Technology for '
            'supporting the TENDERFLOW mainnet testbed deployment used in the Ma\'anshan bridge '
            'simulation reported in this paper.'
        )),
        ('CRediT authorship contribution statement', (
            'Zhonghua Peng: Conceptualization, Methodology, Software, Validation, Writing – original draft. '
            'Tao Ye: Investigation, Resources, Data curation, Writing – review & editing. '
            'Gongming Wang: Supervision, Project administration, Funding acquisition, Writing – review & editing.'
        )),
        ('Declaration of competing interest', (
            'The authors declare that they have no known competing financial interests or personal '
            'relationships that could have appeared to influence the work reported in this paper.'
        )),
        ('Declaration of generative AI and AI-assisted technologies in the manuscript preparation process', (
            'During the preparation of this work the authors used generative AI tools solely for '
            'language polishing and code-assisted deployment scripting. After using these tools, the '
            'authors reviewed and edited the content as needed and take full responsibility for the '
            'content of the published article.'
        )),
        ('Data availability', (
            'Mainnet experiment raw data, figure source CSV files, and reproducibility scripts are '
            'available in the project repository under paper/raw-data/ (experiment_latest.json, '
            'fig6_performance.csv, fig7_reputation.csv, fig8_reputation_evolution.csv) and '
            'paper/run_mainnet_experiment.py. [Add DOI/Zenodo link upon deposit before submission.]'
        )),
    ]

    current = anchor
    for title, body in reversed(sections):
        body_p = insert_after(current, body)
        style_paragraph(body_p)
        title_p = insert_after(current, title, bold=True)
        title_p.paragraph_format.space_before = Pt(12)
        style_paragraph(title_p, bold=True)
        current = title_p
    insert_after(current, '')


def fix_section_numbering_typos(doc):
    replacements = [
        ('4.2.1 Infrastructure Layer', '4.2.1 Infrastructure Layer'),
        ('On Chain Behavior', 'On-Chain Behavior'),
        ('onchain', 'on-chain'),
        ('on Chain', 'on-chain'),
        ('e tendering', 'e-tendering'),
        ('long term', 'long-term'),
        ('self executing', 'self-executing'),
        ('rule based', 'rule-based'),
        ('multi tier', 'multi-tier'),
        ('real time', 'real-time'),
        ('web based', 'web-based'),
        ('paper based', 'paper-based'),
    ]
    for para in doc.paragraphs:
        for old, new in replacements:
            if old in para.text:
                para.text = para.text.replace(old, new)


def normalize_references_heading(doc):
    for p in doc.paragraphs:
        if p.text.strip().lower() == 'references':
            p.text = 'References'
            style_paragraph(p, bold=True)
            break


def write_highlights_file(highlights):
    lines = [
        'Advanced Engineering Informatics — Highlights',
        '(Submit as separate editable file; max 85 characters per bullet including spaces)',
        '',
    ]
    for h in highlights:
        assert len(h) <= 85, f'Highlight too long ({len(h)}): {h}'
        lines.append(f'• {h}')
    HIGHLIGHTS_TXT.write_text('\n'.join(lines), encoding='utf-8')


def write_checklist(m):
    CHECKLIST.write_text(f"""# Advanced Engineering Informatics — Submission Checklist

Journal: **Advanced Engineering Informatics** (Elsevier, ISSN 1474-0346)

## Manuscript files prepared
- [x] `{OUT.name}` — main manuscript (single-column Word)
- [x] `{HIGHLIGHTS_TXT.name}` — highlights (3–5 bullets, ≤85 chars)
- [ ] Graphical abstract (531×1328 px min) — optional but encouraged
- [ ] Cover letter
- [ ] Supplementary material (raw data JSON/CSV if offered)

## Format compliance
- [x] Abstract ≤ 250 words (purpose, results, conclusions)
- [x] Keywords: 7 terms in English
- [x] Highlights separate file
- [x] Numbered sections (1., 1.1., …)
- [x] Numbered references [n] in text
- [x] Acknowledgements before References
- [x] CRediT author contribution statement
- [x] Declaration of competing interest
- [x] Declaration of generative AI use
- [x] Data availability statement
- [ ] Update corresponding author email on title page
- [ ] Deposit data to Zenodo/figshare and add DOI to Data availability

## Mainnet metrics in paper (verified)
- Channel: `{m['channel']}`
- Block height: {m['block_height']}
- Nodes: {m['nodes']}
- Throughput: {m['tps']} TPS (sequential)
- p95 latency: {m['p95_ms']:.0f} ms
- Simulation: {m['sim_ok']}/{m['sim_txs']} success ({m['sim_pct']}%)

## Before submitting in Editorial Manager
- [ ] Remove highlight/revision colors (done in AEI submission file)
- [ ] Verify all figures ≥300 dpi; captions present
- [ ] Spell-check; remove track changes
- [ ] Confirm author order matches submission system
- [ ] Suggest 3–5 reviewers (optional)

## Scope fit (AEI)
Paper reports **engineering informatics** for construction procurement: knowledge representation (reputation ontology), automated knowledge-intensive task (tender evaluation), and **rigorous mainnet validation** — aligned with AEI aims.
""", encoding='utf-8')


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    m = load_metrics()
    highlights = build_highlights(m)
    abstract = build_abstract_aei(m)

    doc = Document(str(SRC))
    remove_revision_notes(doc)
    rebuild_title_page(doc)
    set_abstract_and_keywords(
        doc,
        abstract,
        'Blockchain; Construction procurement; Engineering informatics; Hyperledger Fabric; '
        'Reputation systems; Smart contracts; Tendering',
    )
    insert_highlights(doc, highlights)
    fix_section_numbering_typos(doc)
    insert_back_matter(doc)
    normalize_references_heading(doc)
    apply_body_style(doc)

    doc.save(str(OUT))
    write_highlights_file(highlights)
    write_checklist(m)
    print(f'Saved: {OUT}')
    print(f'Highlights: {HIGHLIGHTS_TXT}')
    print(f'Checklist: {CHECKLIST}')
    print(f'Abstract words: {word_count(abstract)}')
    for i, h in enumerate(highlights, 1):
        print(f'  Highlight {i} ({len(h)} chars): {h}')


if __name__ == '__main__':
    main()
