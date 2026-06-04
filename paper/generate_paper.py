#!/usr/bin/env python3
"""Insert Fig 1–8 at precise Section 4/5/6 positions in the original paper DOCX."""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

FIG_DIR = os.path.join(os.path.dirname(__file__), 'figures')
SRC_DOCX = (
    r'E:\Jp\hyperledge fabric\A Permissioned Enterprise Blockchain Platform for '
    r'Construction Tendering with On-Chain Behavior Reputation_20260602_203742.docx'
)
OUT_DOCX = os.path.join(
    os.path.dirname(__file__),
    'TENDERFLOW_Paper_Mainnet_Data_20260603.docx',
)

EXPERIMENT_JSON = os.path.join(os.path.dirname(__file__), 'raw-data', 'experiment_latest.json')


def load_experiment():
    if not os.path.exists(EXPERIMENT_JSON):
        return {}
    import json
    with open(EXPERIMENT_JSON, encoding='utf-8') as f:
        return json.load(f)

FIGURE_WIDTH = Inches(5.85)


def insert_paragraph_after(paragraph):
    """Return a new empty paragraph inserted immediately after *paragraph*."""
    new_el = paragraph._element.makeelement(qn('w:p'), {})
    paragraph._element.addnext(new_el)
    return Paragraph(new_el, paragraph._parent)


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def add_figure_block(after_paragraph, filename, caption, width=FIGURE_WIDTH):
    """Insert centered figure + caption after the given paragraph. Returns caption paragraph."""
    path = os.path.join(FIG_DIR, filename)
    img_p = insert_paragraph_after(after_paragraph)
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(path):
        img_p.add_run().add_picture(path, width=width)
    else:
        img_p.add_run(f'[Missing figure: {filename}]')

    cap_p = insert_paragraph_after(img_p)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_p.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    return cap_p


def replace_placeholder_with_figure(paragraph, filename, caption):
    """Replace a 'Figure X …' placeholder paragraph with the actual figure block."""
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = os.path.join(FIG_DIR, filename)
    if os.path.exists(path):
        paragraph.add_run().add_picture(path, width=FIGURE_WIDTH)
    else:
        paragraph.add_run(f'[Missing figure: {filename}]')

    # Remove immediately following empty paragraphs (Word split-run artefacts)
    nxt = paragraph._element.getnext()
    while nxt is not None and nxt.tag == qn('w:p'):
        texts = [t.text for t in nxt.iter(qn('w:t')) if t.text and t.text.strip()]
        has_blip = any(el.tag.endswith('blip') for el in nxt.iter())
        if has_blip and not texts:
            to_del = nxt
            nxt = nxt.getnext()
            to_del.getparent().remove(to_del)
            continue
        break

    cap_p = insert_paragraph_after(paragraph)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_p.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    return cap_p


def find_paragraph(doc, predicate):
    for i, p in enumerate(doc.paragraphs):
        if predicate(p):
            return i, p
    return None, None


def find_by_text(doc, fragment, start=0):
    frag = fragment.lower()
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if frag in p.text.lower():
            return i, p
    return None, None


def find_heading(doc, text_fragment):
    return find_by_text(doc, text_fragment)


def insert_network_table(after_paragraph):
    """Table 5 — 14-node configuration (Section 5.2)."""
    exp = load_experiment()
    net = exp.get('network', {})
    cap = insert_paragraph_after(after_paragraph)
    run = cap.add_run("Table 5. 14-Node Network Configuration — Ma'anshan Simulation (Mainnet)")
    run.bold = True
    run.font.size = Pt(11)

    tbl_p = insert_paragraph_after(cap)
    body = insert_paragraph_after(tbl_p)
    rows = [
        ('O1–O5', 'Tender Consensus Cluster', 'RAFT Orderer', '7050–11050'),
        ('R0–R2', 'Provincial Transport Dept.', 'Endorsing Peer (Regulator)', '7051–7055'),
        ('B0–B2', 'Tier-1 Construction Consortium', 'Endorsing Peer (8 Bidders)', '9051–9055'),
        ('A0–A2', 'Audit & Financial Supervision', 'Endorsing Peer (Auditor/Bank)', '11051–11055'),
    ]
    text = (
        'Node ID | Organization Role | Fabric Component | Port\n'
        + '\n'.join(' | '.join(r) for r in rows)
        + f"\n\nChannel: {net.get('channel_id', 'fx-bridge-channel')} | State DB: {net.get('state_database', 'CouchDB')} "
        f"| Chaincode: tenderflow (CCAAS) | Fabric: {net.get('fabric_version', '2.5.12')} "
        f"| Measured block height: {net.get('block_height', 'N/A')} "
        f"| Infrastructure nodes: {net.get('infrastructure_nodes', 14)} "
        f"| Case: Ma'anshan Yangtze River Bridge."
    )
    body.add_run(text).font.size = Pt(9)
    return body


def insert_results_table(after_paragraph):
    """Table 4 — comparative analysis at Section 5.3.3."""
    exp = load_experiment()
    sc = exp.get('scenario_analysis', {})
    bench = exp.get('benchmark', {})
    sim = exp.get('simulation', {})
    cap = insert_paragraph_after(after_paragraph)
    run = cap.add_run('Table 4. Comparative Analysis — Static vs. Dynamic Selection (Mainnet Measured)')
    run.bold = True
    run.font.size = Pt(11)

    body = insert_paragraph_after(cap)
    ok = sum(1 for t in sim.get('transactions', []) if t.get('ok'))
    total = len(sim.get('transactions', [])) or 1
    data = [
        ('Selected Bidder', f"{sc.get('scenario1_static_winner', 'Bidder-A')} (Vi={sc.get('scenario1_static_verified_rep', 0.95)})",
         f"{sc.get('scenario2_dynamic_winner', 'Bidder-A')} (Total={sc.get('scenario2_dynamic_total_score', 0):.3f})"),
        ('High-risk bidder filtered', 'No', f"Yes (Bidder B breach, total={sc.get('bidder_b_total_after_breach', 0):.3f})"),
        ('Simulation success rate', 'N/A', f'{100*ok/total:.1f}% ({ok}/{total} tx)'),
        ('Peak TPS (sequential invoke)', 'N/A', f"{bench.get('tps', 'N/A')}"),
        ('Commit latency (p95)', 'N/A', f"{bench.get('latency_ms', {}).get('p95', 'N/A')} ms"),
        ('Blocks produced (simulation)', 'N/A', str(sim.get('blocks_produced', 'N/A'))),
        ('Final block height', 'N/A', str(exp.get('network', {}).get('block_height', 'N/A'))),
        ('Recovery after breach (Bidder-C)', 'N/A', '12 on-chain Success updates'),
    ]
    lines = ['Metric | Scenario 1 (Static Vi) | Scenario 2 (TENDERFLOW Mainnet)']
    lines += [' | '.join(d) for d in data]
    body.add_run('\n'.join(lines)).font.size = Pt(9)
    return body


# ── Figure placement rules ───────────────────────────────────────────────────
# Section 4: replace existing "Figure X" placeholders
SECTION4_REPLACEMENTS = [
    (
        'figure 2  figure x. schematic architecture',
        'fig2_system_design.png',
        'Figure 2. Schematic architecture of the decentralized tendering system based on '
        'permissioned blockchain and IPFS (TENDERFLOW).',
    ),
    (
        'figure 3.  figure x. interaction',
        'fig3_commit_reveal_workflow.png',
        'Figure 3. Interaction and operational workflow of the decentralized construction '
        'tendering system — Commit-and-Reveal protocol.',
    ),
    (
        'figure 1. the structure of hyperledge fabric',
        'fig1_six_layer_architecture.png',
        'Figure 1. TENDERFLOW six-layer modular architecture for permissioned construction tendering.',
    ),
    (
        'figure 4 .  the diagram of hyperledger fabric',
        'fig4_fabric_network_topology.png',
        'Figure 4. Diagram of the 14-node Hyperledger Fabric network structure deployed for '
        'the Ma\'anshan Yangtze River Bridge simulation.',
    ),
]

# Section 5 & 6: insert after section headings (processed bottom-up to preserve indices)
SECTION_INSERTIONS = [
    (
        '6.3 reputation sensitivity analysis',
        'fig8_reputation_evolution.png',
        'Figure 8. Behavioral reputation evolution demonstrating the asymmetry principle '
        '(η=0.05, θ=0.15; 12 transactions required to recover after a breach).',
        None,
    ),
    (
        '6.1 technical performance benchmarks',
        'fig6_performance_results.png',
        'Figure 6. Technical performance benchmarks: (a) transaction throughput, '
        '(b) transaction latency, (c) qualification verification time.',
        None,
    ),
    (
        '5.3.3 comparative data analysis',
        'fig7_reputation_scenarios.png',
        'Figure 7. Reputation sensitivity analysis — Scenario 1 (qualification-based) vs. '
        'Scenario 2 (TENDERFLOW dynamic behavioral selection).',
        'results_table',
    ),
    (
        '5.2 simulation setup and network configuration',
        'fig5_transaction_sequence.png',
        'Figure 5. Transaction sequence during the Ma\'anshan tendering test '
        '(Commit → Reveal → Reputation Update → Award).',
        'network_table',
    ),
]

# Explorer-captured transaction figures (Section 5.4)
EXPLORER_FIGURES = [
    (
        'fig9_explorer_transaction_feed.png',
        'Figure 9. TENDERFLOWScan explorer — live feed of latest blocks and on-chain tender transactions '
        '(captured from blockchain explorer during simulation).',
    ),
    (
        'fig10_explorer_payload_inspector.png',
        'Figure 10. Read/Write set inspector — decoded tender transaction payload (BID_REVEAL / REPUTATION_UPDATE) '
        'as displayed in the blockchain explorer.',
    ),
    (
        'fig11_explorer_tx_timeline.png',
        'Figure 11. Chronological sequence of on-chain transaction types recorded during the Ma\'anshan test run.',
    ),
]


def apply_section4_replacements(doc):
    replaced = []
    for fragment, filename, caption in SECTION4_REPLACEMENTS:
        _, para = find_by_text(doc, fragment)
        if para is None:
            print(f'  [WARN] Section 4 placeholder not found: {fragment!r}')
            continue
        replace_placeholder_with_figure(para, filename, caption)
        replaced.append(filename)
        print(f'  [OK] Section 4 replaced placeholder → {filename}')
    return replaced


def apply_section56_insertions(doc):
    inserted = []
    for fragment, filename, caption, extra in SECTION_INSERTIONS:
        _, heading = find_by_text(doc, fragment)
        if heading is None:
            print(f'  [WARN] Section heading not found: {fragment!r}')
            continue
        anchor = heading
        if extra == 'network_table':
            anchor = insert_network_table(heading)
        elif extra == 'results_table':
            anchor = insert_results_table(heading)

        add_figure_block(anchor, filename, caption)
        inserted.append(filename)
        print(f'  [OK] Inserted after {fragment!r} → {filename}')
    return inserted


def remove_appendix_if_present(doc):
    """Strip prior 'EXPERIMENTAL IMPLEMENTATION AND FIGURES' appendix from re-runs."""
    to_remove = []
    in_appendix = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == 'EXPERIMENTAL IMPLEMENTATION AND FIGURES':
            in_appendix = True
        if in_appendix:
            to_remove.append(p)
    for p in to_remove:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
    if to_remove:
        print(f'  [OK] Removed {len(to_remove)} appendix paragraph(s) from prior run')


def apply_explorer_figures(doc):
    """Insert Fig 9–11 after Section 5.4 (explorer transaction captures)."""
    _, heading = find_by_text(doc, '5.4 phase-wise implementation')
    if heading is None:
        print('  [WARN] Section 5.4 heading not found — skipping explorer figures')
        return []
    anchor = heading
    inserted = []
    for filename, caption in EXPLORER_FIGURES:
        cap_p = add_figure_block(anchor, filename, caption)
        anchor = cap_p
        inserted.append(filename)
        print(f'  [OK] Explorer figure → {filename}')
    return inserted


def main():
    import generate_figures
    generate_figures.main()

    import generate_explorer_figures
    generate_explorer_figures.main()

    doc = Document(SRC_DOCX)
    print(f'Loaded source: {os.path.basename(SRC_DOCX)}')
    remove_appendix_if_present(doc)

    print('\nSection 4 — replacing Figure placeholders:')
    s4 = apply_section4_replacements(doc)

    print('\nSections 5 & 6 — inserting figures after headings:')
    s56 = apply_section56_insertions(doc)

    print('\nSection 5.4 — explorer transaction figures:')
    s54 = apply_explorer_figures(doc)

    doc.save(OUT_DOCX)
    print(f'\nSaved: {OUT_DOCX}')

    import full_paper_mainnet_revision
    full_paper_mainnet_revision.SRC = Path(OUT_DOCX)
    full_paper_mainnet_revision.main()
    print(f'Figures placed: {len(s4) + len(s56) + len(s54)} total')


if __name__ == '__main__':
    main()
